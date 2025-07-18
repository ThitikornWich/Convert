import streamlit as st
import pandas as pd
import re
import io
from docx import Document

def extract_sections_from_doc(doc_file):
    # อ่านไฟล์ Word
    doc = Document(doc_file)
    # ดึงข้อความทุก cell ของทุกตาราง + ทุก paragraph (ไม่ตกหล่น)
    all_lines = []
    # ตาราง
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    all_lines.extend([t for t in text.split('\n') if t.strip()])
    # ย่อหน้า
    for para in doc.paragraphs:
        t = para.text.strip()
        if t:
            all_lines.append(t)

    # รวมเนื้อหาทุกบรรทัดที่ clean แล้ว
    lines = [l.strip() for l in all_lines if l.strip()]
    
    # ฟังก์ชันช่วยหา value ที่ตามหลัง label (case-insensitive)
    def extract_after_label(label):
        for i, line in enumerate(lines):
            if re.search(rf'^{label}\b', line, re.IGNORECASE):
                # เช่น Title: xxx
                m = re.match(rf'{label}[:：\s]*(.+)', line, re.IGNORECASE)
                if m:
                    return m.group(1).strip()
                # หรือ Title [ขึ้นบรรทัดใหม่]
                elif i+1 < len(lines):
                    return lines[i+1].strip()
        return ""
    
    # Title
    title = extract_after_label("Title")
    # Description
    description = extract_after_label("Description")
    # H1 (Heading 1)
    h1 = extract_after_label("H1")
    if not h1:
        h1 = extract_after_label("Heading 1")
    
    # Product Section
    product_section = ""
    # pattern รองรับทั้ง ------- Product Section --------- หรือ Product Section เฉยๆ
    for i, line in enumerate(lines):
        if re.search(r'(?:-{3,}\s*)?Product\s*Section', line, re.IGNORECASE):
            # เอาตั้งแต่บรรทัดถัดไปเรื่อย ๆ จนกว่าจะเจอหัวข้อใหม่หรือจบ
            j = i + 1
            part = []
            while j < len(lines):
                # ถ้าเจอ label ใหม่ (Title, Description, H1, Heading 1, Slug, หรือ Product Section อีก)
                if re.search(r'^(Title|Description|H1|Heading 1|Slug|Product\s*Section)\b', lines[j], re.IGNORECASE):
                    break
                part.append(lines[j])
                j += 1
            product_section = "\n".join(part).strip()
            break

    return {
        'Title': title,
        'Description': description,
        'Heading 1': h1,
        'Product Section': product_section
    }

# ========== Streamlit UI ==========
st.title("Word Content Extractor (Title, Description, H1, Product Section)")
st.write("อัปโหลดไฟล์ Word (.docx) หลายไฟล์พร้อมกันเพื่อดึงข้อมูลและ export เป็น CSV/XLSX")

uploaded_files = st.file_uploader(
    "อัปโหลด Word Documents (.docx)", 
    type=["docx"], 
    accept_multiple_files=True,
    help="สามารถเลือกหลายไฟล์พร้อมกันได้"
)

if uploaded_files:
    st.info(f"อัปโหลดไฟล์สำเร็จ: {len(uploaded_files)} ไฟล์")
    
    # แสดงรายชื่อไฟล์ที่อัปโหลด
    with st.expander(f"📁 รายชื่อไฟล์ที่อัปโหลด ({len(uploaded_files)} ไฟล์)"):
        for i, file in enumerate(uploaded_files, 1):
            st.write(f"{i}. {file.name} ({file.size:,} bytes)")
    
    if st.button("🚀 Start Process All Files", type="primary"):
        progress_bar = st.progress(0)
        status_text = st.empty()
        
        all_results = []
        
        for i, uploaded_file in enumerate(uploaded_files):
            status_text.text(f"กำลังประมวลผล: {uploaded_file.name} ({i+1}/{len(uploaded_files)})")
            progress_bar.progress((i + 1) / len(uploaded_files))
            
            try:
                extracted = extract_sections_from_doc(uploaded_file)
                # เพิ่มชื่อไฟล์เป็นคอลัมน์แรก
                result = {
                    'Filename': uploaded_file.name,
                    'Title': extracted['Title'],
                    'Description': extracted['Description'],
                    'Heading 1': extracted['Heading 1'],
                    'Product Section': extracted['Product Section']
                }
                all_results.append(result)
            except Exception as e:
                st.error(f"เกิดข้อผิดพลาดในไฟล์ {uploaded_file.name}: {str(e)}")
                # เพิ่มข้อมูล error
                result = {
                    'Filename': uploaded_file.name,
                    'Title': f'Error: {str(e)}',
                    'Description': '',
                    'Heading 1': '',
                    'Product Section': ''
                }
                all_results.append(result)
        
        status_text.text("✅ ประมวลผลเสร็จสิ้น!")
        
        if all_results:
            st.success(f"ประมวลผลไฟล์เสร็จสิ้น: {len(all_results)} ไฟล์")
            
            # แสดงผลลัพธ์
            df = pd.DataFrame(all_results)
            
            st.subheader("📋 ผลลัพธ์การดึงข้อมูล:")
            st.dataframe(df, use_container_width=True)
            
            # แสดงสถิติ
            st.subheader("📊 สถิติการประมวลผล:")
            col1, col2, col3, col4 = st.columns(4)
            
            with col1:
                successful_files = len([r for r in all_results if not r['Title'].startswith('Error:')])
                st.metric("ไฟล์สำเร็จ", successful_files, f"จาก {len(all_results)}")
            
            with col2:
                title_count = len([r for r in all_results if r['Title'] and not r['Title'].startswith('Error:')])
                st.metric("มี Title", title_count)
            
            with col3:
                desc_count = len([r for r in all_results if r['Description']])
                st.metric("มี Description", desc_count)
            
            with col4:
                product_count = len([r for r in all_results if r['Product Section']])
                st.metric("มี Product Section", product_count)
            
            # ปุ่ม Export
            st.subheader("💾 Export ข้อมูล:")
            col1, col2 = st.columns(2)
            
            with col1:
                # Export CSV
                csv_buffer = io.StringIO()
                df.to_csv(csv_buffer, index=False, encoding='utf-8-sig')
                st.download_button(
                    "📥 ดาวน์โหลดเป็น CSV", 
                    csv_buffer.getvalue(), 
                    file_name=f"extracted_content_{len(uploaded_files)}_files.csv",
                    mime="text/csv",
                    use_container_width=True
                )
            
            with col2:
                # Export XLSX
                xlsx_buffer = io.BytesIO()
                with pd.ExcelWriter(xlsx_buffer, engine='xlsxwriter') as writer:
                    df.to_excel(writer, index=False, sheet_name='Extracted Content')
                    # ปรับขนาดคอลัมน์
                    worksheet = writer.sheets['Extracted Content']
                    for idx, col in enumerate(df.columns):
                        max_length = max(
                            df[col].astype(str).apply(len).max(),
                            len(col)
                        )
                        worksheet.set_column(idx, idx, min(max_length + 2, 50))
                
                st.download_button(
                    "📥 ดาวน์โหลดเป็น XLSX", 
                    xlsx_buffer.getvalue(), 
                    file_name=f"extracted_content_{len(uploaded_files)}_files.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    use_container_width=True
                )
            
            # แสดงรายละเอียดแต่ละไฟล์
            st.subheader("📄 รายละเอียดแต่ละไฟล์:")
            for i, result in enumerate(all_results, 1):
                with st.expander(f"📄 {i}. {result['Filename']}"):
                    if result['Title'].startswith('Error:'):
                        st.error(f"เกิดข้อผิดพลาด: {result['Title']}")
                    else:
                        col1, col2 = st.columns(2)
                        with col1:
                            st.write("**Title:**")
                            st.text_area("", result['Title'], height=60, key=f"title_{i}")
                            st.write("**Heading 1:**")
                            st.text_area("", result['Heading 1'], height=60, key=f"h1_{i}")
                        
                        with col2:
                            st.write("**Description:**")
                            st.text_area("", result['Description'], height=60, key=f"desc_{i}")
                            st.write("**Product Section:**")
                            st.text_area("", result['Product Section'], height=100, key=f"product_{i}")
else:
    st.info("กรุณาอัปโหลดไฟล์ Word Document (.docx) เพื่อเริ่มต้น")
    
    # เพิ่มคำแนะนำ
    with st.expander("📋 คำแนะนำการใช้งาน"):
        st.markdown("""
        **การอัปโหลดหลายไฟล์:**
        1. คลิกปุ่ม "Browse files" 
        2. เลือกไฟล์ Word (.docx) หลายไฟล์พร้อมกัน (ใช้ Ctrl+Click หรือ Shift+Click)
        3. กด "Start Process All Files" เพื่อประมวลผลทุกไฟล์
        
        **รูปแบบเอกสารที่รองรับ:**
        - Title: ต้องมีคำว่า "Title" ในเอกสาร
        - Description: ต้องมีคำว่า "Description" ในเอกสาร  
        - H1: ต้องมีคำว่า "H1" หรือ "Heading 1"
        - Product Section: ต้องมีคำว่า "Product Section" (เนื้อหาด้านล่างทั้งหมดจะถูกดึงมา)
        
        **ผลลัพธ์:**
        - ไฟล์ CSV/XLSX จะรวมข้อมูลจากทุกไฟล์ในตารางเดียว
        - คอลัมน์แรกจะเป็นชื่อไฟล์เพื่อแยกแยะ
        """)

